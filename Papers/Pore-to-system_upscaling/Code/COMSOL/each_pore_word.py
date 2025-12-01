import numpy as np
from docx import Document
from pathlib import Path
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from tqdm import tqdm,trange
Path_root = Path(r'D:\yjp\Workdir\Code\ZJU\Study\Python\multi-physic-network-model\Samples\heat_transfer_Finney\Pore_Distribution')
doc = Document()
num_void = 1854
num_solid = 1559
num_pore = num_void + num_solid

for i in trange(num_pore):
    para_i=doc.add_paragraph()
    para_i.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_i.add_run().add_picture(f"{Path_root}/T_distribution_pore_{i}.png", height=Cm(3.6))
    para_i.add_run().add_picture(f"{Path_root}/U_distribution_pore_{i}.png", height=Cm(3.6))
    para_i.add_run().add_picture(f"{Path_root}/P_distribution_pore_{i}.png", height=Cm(3.6))


doc.save(f"{Path_root}/Pore_Distribution.docx")